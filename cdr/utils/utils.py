#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""=================================================
    @Project: django_hip_service
    @File： utils.py
    @Author：liaozhimingandy
    @Email: liaozhimingandy@gmail.com
    @Date：2024-12-02 15:36
    @Desc:
================================================="""
from django.apps import apps
from docx import Document


def get_choices_display(field):
    """获取字段的枚举选项"""
    if field.choices:
        choices_list = [f"{value}: {label}" for value, label in field.choices]
        return "\n".join(choices_list)
    return "N/A"


def export_models_to_word(app_label, output_file, exclude_fields=None):
    # 创建 Word 文档
    doc = Document()
    doc.add_heading(f"Models in {app_label.capitalize()} App", level=1)

    # 获取应用中的所有模型
    app_config = apps.get_app_config(app_label)
    models = app_config.get_models()

    # 排除字段列表，默认不排除任何字段
    exclude_fields = exclude_fields or []

    for model in models:

        # 获取模型的英文和中文名称
        model_name_en = model.__name__  # 英文名称（模型类名）
        model_name_cn = model._meta.verbose_name  # 中文名称（verbose_name）

        # 添加模型名称（英文 + 中文）
        doc.add_heading(f"{model_name_en} ({model_name_cn})", level=2)

        # 创建字段表
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "字段名称"
        hdr_cells[1].text = "描述"
        hdr_cells[2].text = "类型"
        hdr_cells[3].text = "枚举值"

        # 填充字段信息
        for field in model._meta.fields:
            # 如果字段在排除列表中，则跳过
            if field.name in exclude_fields:
                continue

            row_cells = table.add_row().cells
            row_cells[0].text = field.name
            row_cells[1].text = field.verbose_name or "N/A"
            row_cells[2].text = field.get_internal_type()
            row_cells[3].text = get_choices_display(field)

    # 保存文档
    doc.save(output_file)
    print(f"文档已保存到 {output_file}")


if __name__ == "__main__":
    # 调用函数
    # 用法,使用python控制台运行
    # utils.export_models_to_word('cdr', 'D:\\models.docx', exclude_fields=["id", "gmt_created", "from_src"])
    export_models_to_word("cdr", "D:\\models.docx")
